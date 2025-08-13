import os
from flask import Flask, render_template, request, Response, send_file, jsonify, redirect, url_for
import csv
from io import BytesIO
from pypxlib import Table
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3

app = Flask(__name__)

# Database setup for persistent groups
DATABASE = 'app.db'

def init_db():
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS groups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE NOT NULL
            )
        ''')
        c.execute('''
            CREATE TABLE IF NOT EXISTS group_items (
                group_id INTEGER NOT NULL,
                item_id INTEGER NOT NULL,
                FOREIGN KEY(group_id) REFERENCES groups(id),
                UNIQUE(group_id, item_id)
            )
        ''')
        c.execute('''
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        ''')
        # Initialize default settings if not exists
        c.execute('''
            INSERT OR IGNORE INTO settings (key, value) 
            VALUES ('db_path', 'items.DB')
        ''')
        conn.commit()

def get_setting(key):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('SELECT value FROM settings WHERE key = ?', (key,))
        row = c.fetchone()
        return row[0] if row else None

def set_setting(key, value):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('''
            INSERT INTO settings (key, value) VALUES (?, ?)
            ON CONFLICT(key) DO UPDATE SET value = excluded.value
        ''', (key, value))
        conn.commit()

def get_groups():
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('SELECT id, name FROM groups ORDER BY name')
        return [{'id': row[0], 'name': row[1]} for row in c.fetchall()]

def add_group(name):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        try:
            c.execute('INSERT INTO groups (name) VALUES (?)', (name,))
            conn.commit()
            return c.lastrowid
        except sqlite3.IntegrityError:
            return None  # Group already exists

def update_group(group_id, new_name):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        try:
            c.execute('UPDATE groups SET name = ? WHERE id = ?', (new_name, group_id))
            conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False  # Group name already exists

def delete_group(group_id):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('DELETE FROM group_items WHERE group_id = ?', (group_id,))
        c.execute('DELETE FROM groups WHERE id = ?', (group_id,))
        conn.commit()

def add_item_to_group(group_id, item_id):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        try:
            c.execute('INSERT INTO group_items (group_id, item_id) VALUES (?, ?)', (group_id, item_id))
            conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False  # Item already in group

def remove_item_from_group(group_id, item_id):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('DELETE FROM group_items WHERE group_id = ? AND item_id = ?', (group_id, item_id))
        conn.commit()

def get_group_items(group_id):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('SELECT item_id FROM group_items WHERE group_id = ?', (group_id,))
        return [row[0] for row in c.fetchall()]

# Cache the table data to improve performance
TABLE_DATA = None
LAST_REFRESH = 0
CACHE_TIMEOUT = 300  # 5 minutes

def get_records():
    """Read records from Paradox table with caching"""
    global TABLE_DATA, LAST_REFRESH
    # Return cached data if it's still fresh
    if TABLE_DATA and time.time() - LAST_REFRESH < CACHE_TIMEOUT:
        return TABLE_DATA
    
    # Get DB path from settings
    db_path = get_setting('db_path') or 'items.DB'
    
    try:
        table = Table(db_path, encoding='windows-1251')
        records = []
        # Extract field names from the table schema
        field_names = [field for field in table.fields]
        # Convert records to dictionaries for easier handling
        for record in table:
            record_dict = {field: getattr(record, field) for field in field_names}
            records.append(record_dict)
        TABLE_DATA = (records, field_names)
        LAST_REFRESH = time.time()
        return TABLE_DATA
    except Exception as e:
        print(f"Error reading table: {e}")
        return [], []

@app.route('/')
def home():
    """Home page with options"""
    return render_template('home.html')

@app.route('/manage_groups')
def manage_groups():
    """Group management page"""
    return render_template('manage_groups.html')

@app.route('/edit_group/<int:group_id>')
def edit_group(group_id):
    """Edit group page"""
    return render_template('edit_group.html', group_id=group_id)

@app.route('/settings')
def settings_page():
    """Settings page"""
    db_path = get_setting('db_path') or ''
    return render_template('settings.html', db_path=db_path)

@app.route('/search', methods=['GET'])
def search_items():
    """Server-side search endpoint"""
    query = request.args.get('q', '').lower()
    page = int(request.args.get('page', 1))
    per_page = 50
    records, field_names = get_records()

    # Filter records based on query
    filtered = []
    if query:
        for record in records:
            # Create searchable text from relevant fields
            search_text = f"{record.get('id', '')} {record.get('Code', '')} {record.get('Item', '')} {record.get('ClientPrice', '')} {record.get('Vendor', '')}".lower()
            if query in search_text:
                filtered.append(record)
    else:
        # Return all records if no query
        filtered = records

    # Pagination
    total = len(filtered)
    start = (page - 1) * per_page
    end = start + per_page
    paginated = filtered[start:end]

    # Prepare simplified data for display
    results = []
    for record in paginated:
        results.append({
            'id': record.get('id', ''),
            'Code': record.get('Code', ''),
            'Item': record.get('Item', ''),
            'ClientPrice': record.get('ClientPrice', ''),
            'Vendor': record.get('Vendor', ''),
            'VendorPrice': record.get('VendorPrice', '')
        })

    return jsonify({
        'results': results,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page
    })

@app.route('/groups', methods=['GET'])
def api_get_groups():
    return jsonify(get_groups())

@app.route('/groups', methods=['POST'])
def create_group():
    name = request.json.get('name')
    if not name:
        return jsonify({'error': 'Group name is required'}), 400
    group_id = add_group(name)
    if group_id:
        return jsonify({'id': group_id, 'name': name}), 201
    else:
        return jsonify({'error': 'Group already exists'}), 400

@app.route('/groups/<int:group_id>', methods=['GET'])
def get_group(group_id):
    with sqlite3.connect(DATABASE) as conn:
        c = conn.cursor()
        c.execute('SELECT id, name FROM groups WHERE id = ?', (group_id,))
        row = c.fetchone()
        if row:
            return jsonify({'id': row[0], 'name': row[1]})
        else:
            return jsonify({'error': 'Group not found'}), 404

@app.route('/groups/<int:group_id>', methods=['PUT'])
def update_group_api(group_id):
    new_name = request.json.get('name')
    if not new_name:
        return jsonify({'error': 'Group name is required'}), 400
    if update_group(group_id, new_name):
        return jsonify({'message': 'Group updated'}), 200
    else:
        return jsonify({'error': 'Group name already exists'}), 400

@app.route('/groups/<int:group_id>', methods=['DELETE'])
def delete_group_api(group_id):
    delete_group(group_id)
    return jsonify({'message': 'Group deleted'}), 200

@app.route('/groups/<int:group_id>/items', methods=['GET'])
def get_group_items_api(group_id):
    item_ids = get_group_items(group_id)
    return jsonify(item_ids)

@app.route('/groups/items', methods=['POST'])
def add_item_to_group_api():
    data = request.json
    group_id = data.get('group_id')
    item_id = data.get('item_id')
    if not group_id or not item_id:
        return jsonify({'error': 'Missing group_id or item_id'}), 400
    if add_item_to_group(group_id, item_id):
        return jsonify({'message': 'Item added to group'}), 201
    else:
        return jsonify({'error': 'Item already in group'}), 400

@app.route('/groups/<int:group_id>/items/<int:item_id>', methods=['DELETE'])
def remove_item_from_group_api(group_id, item_id):
    remove_item_from_group(group_id, item_id)
    return jsonify({'message': 'Item removed from group'}), 200

@app.route('/settings', methods=['GET'])
def api_get_settings():
    db_path = get_setting('db_path') or ''
    return jsonify({'db_path': db_path})

@app.route('/settings', methods=['POST'])
def api_update_settings():
    db_path = request.json.get('db_path')
    if db_path is None:
        return jsonify({'error': 'Missing db_path parameter'}), 400
    
    set_setting('db_path', db_path)
    
    # Clear cache to force reload with new path
    global TABLE_DATA, LAST_REFRESH
    TABLE_DATA = None
    LAST_REFRESH = 0
    
    return jsonify({'message': 'Settings updated'}), 200

@app.route('/export_word', methods=['GET'])
def export_word():
    """Export all groups as Word document"""
    # Get all groups
    groups_data = get_groups()
    group_ids = [group['id'] for group in groups_data]

    if not group_ids:
        return "No groups found", 404

    # Get all records
    records, field_names = get_records()
    records_dict = {str(r['id']): r for r in records}

    # Create Word document
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Add title
    title = doc.add_paragraph('НАЦИОНАЛНА АГЕНЦИЯ ЗА ПРИХОДИТЕ')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(14)
    title_run.font.bold = True

    # Add subtitle
    subtitle = doc.add_paragraph('ЦЕНТРАЛНО УПРАВЛЕНИЕ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.bold = True

    # Add second subtitle
    subtitle2 = doc.add_paragraph('ГЛАВНА ДИРЕКЦИЯ “ФИСКАЛЕН КОНТРОЛ“')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(12)
    subtitle2.runs[0].font.bold = True

    # Add horizontal line
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('________________________________________________')

    # Add address
    address = doc.add_paragraph('1000 София. бул. “Княз Дондуков“ №52 Телефон: 0700 18 700 Факс: (02) 9859 3099')
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add appendix
    doc.add_paragraph('Приложение №1 към Протокол №……………………………..')

    # Add company info
    doc.add_paragraph('Задължено лице: Анет4 ЕООД')
    doc.add_paragraph('ЕИК: 202112929')
    doc.add_paragraph('Търговски обект: ? от Анет4 гр. Бургас ул. ? А93')

    # Add date
    today = datetime.now().strftime('%d.%m.%Y')
    doc.add_paragraph(f'Данни за цените на продуктите към дата: {today}')

    # Add a space
    doc.add_paragraph()

    # Add tables for each group
    for group_idx, group in enumerate(groups_data):
        # Get items for this group
        item_ids = get_group_items(group['id'])
        if not item_ids:
            continue

        # Add group header
        p = doc.add_paragraph(f'{group_idx+1}. {group["name"]}')
        p.runs[0].font.bold = True

        # Create table for items
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Set column widths
        col_widths = [Cm(1.5), Cm(8), Cm(3), Cm(3)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        # Add header row
        hdr_cells = table.rows[0].cells
        headers = ['№', 'Марка', 'Продажна цена с ДДС', 'Доставна цена без ДДС']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add items to the table
        for item_idx, item_id in enumerate(item_ids):
            item = records_dict.get(str(item_id))
            if not item:
                continue
            row_cells = table.add_row().cells
            # Item numbering
            row_cells[0].text = f'{group_idx+1}.{item_idx+1}'
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Item name
            row_cells[1].text = item.get('Item', '')
            # Format prices
            try:
                client_price = float(item.get('ClientPrice', '0.0000'))
                vendor_price = float(item.get('VendorPrice', '0.0000'))
            except ValueError:
                client_price = 0.0
                vendor_price = 0.0
            # Client price with VAT
            row_cells[2].text = f"{client_price:.4f}"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Vendor price without VAT
            row_cells[3].text = f"{vendor_price:.4f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add space after table
        doc.add_paragraph()

    # Add footer - simplified version
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Create footer paragraph
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.add_run("ЦУ на НАП 2025г").font.size = Pt(9)

    # Save document to in-memory file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Return as downloadable Word file
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    init_db()
    app.run(debug=True)
